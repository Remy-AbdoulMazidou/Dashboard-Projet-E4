"""
Génère docs/rapport_dashboard_remy.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = "docs/rapport_dashboard_remy.docx"
os.makedirs("docs", exist_ok=True)

doc = Document()

# ── styles helpers ────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    return p

def bold_body(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + " : ")
    set_font(r, bold=True)
    r2 = p.add_run(text)
    set_font(r2)

def placeholder():
    p = doc.add_paragraph()
    r = p.add_run("[ Insérer ici une capture d'écran du dashboard ]")
    set_font(r, size=10, color=(150, 150, 150))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)

# ── PAGE DE TITRE ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rapport de mi-projet — Ma contribution")
set_font(r, bold=True, size=20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Dashboard analytique — Visualisation des données du projet E4")
set_font(r2, size=14, color=(46, 116, 181))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Rémy ABDOUL MAZIDOU — ESIEE Paris — 2024-2025")
set_font(r3, size=11, color=(100, 100, 100))

doc.add_page_break()

# ── 1. MA CONTRIBUTION ────────────────────────────────────────────────────────
h1("1. Ma contribution au projet")

body(
    "Rémy a développé un outil de visualisation interactif sous forme de dashboard web, "
    "permettant d'explorer et de comparer les données produites par les différentes parties du projet. "
    "L'objectif est de rendre ces données lisibles et accessibles sans avoir à écrire de code : "
    "une fois les fichiers CSV générés par l'algorithme d'analyse d'images, le dashboard permet "
    "de les visualiser immédiatement sous forme de graphiques interactifs."
)
body(
    "Pour l'instant, les données utilisées sont simulées à partir de scripts de génération, "
    "afin de pouvoir développer et tester l'outil en amont de la réception des vraies données "
    "issues de l'algorithme."
)

# ── 2. POURQUOI UN DASHBOARD ──────────────────────────────────────────────────
h1("2. Pourquoi faire un dashboard analytique ?")

body(
    "Dans un projet combinant analyse d'images, simulation numérique et apprentissage automatique, "
    "les données produites par chaque partie sont nombreuses et hétérogènes. Un dashboard analytique "
    "permet de les centraliser et de les rendre exploitables dans une interface unique, accessible à "
    "tous les membres du groupe sans nécessiter de compétences en programmation."
)
body(
    "L'intérêt principal réside dans la capacité à visualiser rapidement des tendances, comparer des "
    "matériaux sur plusieurs indicateurs simultanément et détecter des comportements inhabituels. "
    "Dans le cadre de ce projet, où l'objectif est de relier des paramètres microstructuraux à des "
    "propriétés acoustiques, la visualisation constitue une étape analytique à part entière : elle "
    "permet de formuler des hypothèses avant toute modélisation et d'orienter les choix pour la "
    "suite du travail."
)
body(
    "Le dashboard est également un outil de communication utile pour présenter les résultats de "
    "façon claire et accessible, que ce soit au sein du groupe, lors des soutenances, ou avec les "
    "partenaires du laboratoire."
)

# ── 3. CHOIX TECHNIQUES ───────────────────────────────────────────────────────
h1("3. Choix techniques")

body(
    "Le dashboard a été développé sur Visual Studio Code en langage Python avec le framework "
    "Plotly Dash, qui permet de créer des applications web interactives sans écrire de JavaScript. "
    "Le code est organisé de façon modulaire : chaque fichier correspond à un rôle précis dans "
    "l'application."
)

bold_body("app.py", "point d'entrée de l'application. Initialise le serveur Dash, assemble les "
    "onglets et définit les callbacks qui mettent à jour les graphiques en fonction des actions "
    "de l'utilisateur.")
bold_body("config.py", "regroupe toutes les constantes visuelles : palette de couleurs, thème, "
    "mise en forme des graphiques. Centraliser ces valeurs permet de modifier l'apparence du "
    "dashboard en un seul endroit.")
bold_body("data.py", "gère le chargement et la préparation des données. Lit les fichiers CSV, "
    "fusionne les métadonnées et expose des fonctions utilitaires pour filtrer les données "
    "selon le matériau ou l'échantillon sélectionné.")
bold_body("components.py", "contient les composants réutilisables de l'interface : cartes "
    "graphiques, bannières explicatives et indicateurs clés. Chaque carte gère ses propres "
    "filtres de matériaux indépendamment des autres.")
body(
    "Les quatre onglets sont chacun développés dans un fichier séparé (tab_overview.py, "
    "tab_morphology.py, tab_acoustics.py, tab_comparison.py), exposant une fonction de mise "
    "en page et une ou plusieurs fonctions de construction des figures Plotly."
)

# ── 4. PRÉSENTATION DU DASHBOARD ──────────────────────────────────────────────
h1("4. Présentation du dashboard")

body(
    "Le dashboard est organisé en quatre onglets, chacun correspondant à un niveau d'analyse "
    "différent des données du projet."
)

h2("Vue d'ensemble")
body(
    "Page d'accueil du dashboard. Elle affiche les indicateurs clés du jeu de données "
    "(nombre d'échantillons, fibres détectées, contacts entre fibres, porosité moyenne) "
    "ainsi qu'un tableau récapitulatif qui compare l'ensemble des matériaux sur leurs métriques "
    "principales : diamètre moyen, longueur, porosité et absorption acoustique."
)
placeholder()

h2("Morphologie des fibres")
body(
    "Cet onglet permet d'analyser la distribution des diamètres de fibres pour chaque matériau, "
    "à travers un boxplot comparatif et une courbe de densité. Ces deux graphiques permettent "
    "de voir si les fibres d'un matériau sont homogènes en taille ou au contraire très dispersées, "
    "ce qui, d'après les travaux de Tran et al. (2024), a un impact direct sur les propriétés "
    "acoustiques."
)
placeholder()

h2("Propriétés acoustiques")
body(
    "Cet onglet présente les courbes d'absorption sonore de chaque échantillon sur cinq fréquences "
    "allant de 250 Hz à 4 kHz, ainsi qu'un graphique croisant la porosité et la résistivité à l'air "
    "des matériaux, avec une courbe de tendance ajustée automatiquement."
)
placeholder()

h2("Comparaison")
body(
    "C'est l'onglet le plus directement lié à l'objectif du projet. Un graphique à barres compare "
    "les matériaux fréquence par fréquence, et un scatter plot met en relation le diamètre moyen "
    "des fibres et l'absorption à 1 kHz. Ce dernier graphique permet de tester visuellement si "
    "les fibres fines absorbent mieux le son que les fibres épaisses, et d'orienter la modélisation "
    "en conséquence."
)
placeholder()

# ── 5. LIEN AVEC LES ARTICLES ─────────────────────────────────────────────────
h1("5. Lien avec les articles scientifiques")

body(
    "Deux articles ont été utilisés comme base pour orienter les choix de visualisation et "
    "comprendre les paramètres clés à analyser."
)
bold_body(
    "Depriester et al. (2022)",
    "décrit l'algorithme d'extraction des caractéristiques géométriques des fibres à partir "
    "d'images de microtomographie X. Cet article a servi de référence pour comprendre la "
    "structure des données produites (diamètre, orientation, longueur par fibre) et donc "
    "décider quelles variables visualiser en priorité dans le dashboard."
)
bold_body(
    "Tran et al. (2024)",
    "étudie l'impact de la polydispersité des fibres sur les propriétés acoustiques. "
    "La courbe de densité (KDE) de l'onglet Morphologie est directement inspirée de la "
    "Figure 3b de cet article, qui montre comment la dispersion des diamètres influence "
    "le coefficient d'absorption. Le scatter plot diamètre/absorption de l'onglet Comparaison "
    "est également construit pour tester visuellement l'hypothèse centrale de cet article."
)

# ── 6. AMÉLIORATIONS ET SUITE ─────────────────────────────────────────────────
h1("6. Quelles améliorations ? La suite ?")

body(
    "Le dashboard est fonctionnel et prêt à recevoir les vraies données dès qu'elles seront "
    "disponibles. La première étape sera d'intégrer les sorties réelles de l'algorithme d'analyse "
    "d'images et de vérifier que les graphiques restent cohérents avec des valeurs réelles."
)
body(
    "Ensuite, de nouveaux graphiques pourront être ajoutés en fonction des besoins qui émergeront "
    "lors de l'analyse des résultats. L'outil est conçu pour évoluer facilement au fil du projet."
)

doc.save(OUT)
print(f"Rapport créé : {OUT}")
